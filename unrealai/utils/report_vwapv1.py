import os
import zipfile
import datetime as dt
from dateutil.relativedelta import relativedelta

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.lines import Line2D
from matplotlib.backends.backend_pdf import PdfPages

import mplfinance as mpf
import norgatedata
from scipy.signal import argrelextrema
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

# ────────────────────────── CONFIG / STYLE ────────────────────────────
mc = mpf.make_marketcolors(up="green", down="red", edge="inherit", wick="inherit", volume="inherit")
style_dark = mpf.make_mpf_style(base_mpf_style="nightclouds", marketcolors=mc)

# ────────────────────────── DATA HELPERS ──────────────────────────────

def fetch_10yr_daily(symbol: str, yrs: int = 10, end = None) -> pd.DataFrame | None:
    if end:
        end = pd.to_datetime(end)
        start = end - relativedelta(years=yrs)
    else:
        end = dt.datetime.today()
        start = end - relativedelta(years=yrs)
    try:
        return norgatedata.price_timeseries(
            symbol,
            stock_price_adjustment_setting=norgatedata.StockPriceAdjustmentType.TOTALRETURN,
            padding_setting=norgatedata.PaddingType.NONE,
            start_date=pd.Timestamp(start.date()),
            end_date = pd.Timestamp(end.date()),
            timeseriesformat="pandas-dataframe",
        )
    except Exception as exc:
        print(f"[{symbol}] Norgate error → {exc}")
        return None


def resample_ohlc(df: pd.DataFrame, freq: str) -> pd.DataFrame:
    return (
        df.resample(freq)
        .agg({"Open": "first", "High": "max", "Low": "min", "Close": "last", "Volume": "sum"})
        .dropna()
    )


def calc_vwap(df: pd.DataFrame) -> None:
    df["adj_close"] = df["Close"]
    hi = argrelextrema(df["adj_close"].values, np.greater, order=5)[0]
    lo = argrelextrema(df["adj_close"].values, np.less, order=5)[0]
    if len(hi) == 0 or len(lo) == 0:
        df[["vwap_maxloc", "vwap_minloc", "vwap_highvol"]] = np.nan
        return
    maxloc = hi[np.argmax(df["adj_close"].iloc[hi])]
    minloc = lo[np.argmin(df["adj_close"].iloc[lo])]
    hv_idx = df.index.get_loc(df["Volume"].iloc[125:].idxmax()) if len(df) > 125 else 0

    for label, start in (
        ("vwap_maxloc", maxloc),
        ("vwap_minloc", minloc),
        ("vwap_highvol", hv_idx),
    ):
        pxv = (df["adj_close"] * df["Volume"]).iloc[start:]
        vol = df["Volume"].iloc[start:]
        df[label] = pxv.expanding().sum() / vol.expanding().sum()

# ────────────────────────── PLOT HELPERS ──────────────────────────────

def _beautify(ax):
    ax.yaxis.tick_right(); ax.yaxis.set_label_position("right"); ax.grid(alpha=0.2)


def build_fig(symbol: str, end = None):
    df_d = fetch_10yr_daily(symbol, end=end)
    if df_d is None or df_d.empty:
        return None

    df_m, df_w = resample_ohlc(df_d, "M"), resample_ohlc(df_d, "W")
    df_1y = df_d.iloc[-252:].copy(); calc_vwap(df_1y)

    jan1 = pd.Timestamp(f"{dt.datetime.today().year}-01-01")
    slice_ytd = df_d[df_d.index >= jan1]
    ytd_str = f"{((slice_ytd['Close'].iat[-1]/slice_ytd['Close'].iat[0]-1)*100):+.2f}%" if len(slice_ytd)>=2 else "N/A"

    fig = mpf.figure(figsize=(14,10), style=style_dark, facecolor="black")
    ax_m, ax_w, ax_d = fig.add_subplot(2,2,1), fig.add_subplot(2,2,2), fig.add_subplot(2,1,2)

    mpf.plot(df_m, type="candle", ax=ax_m, style=style_dark, volume=False, show_nontrading=True)
    mpf.plot(df_w, type="candle", ax=ax_w, style=style_dark, volume=False, show_nontrading=True)

    addp=[mpf.make_addplot(df_1y['vwap_maxloc'],ax=ax_d,color='lightblue'),
          mpf.make_addplot(df_1y['vwap_minloc'],ax=ax_d,color='pink'),
          mpf.make_addplot(df_1y['vwap_highvol'],ax=ax_d,color='yellow')]
    mpf.plot(df_1y, type='candle', ax=ax_d, addplot=addp, style=style_dark, volume=False, show_nontrading=True)

    last_close=df_1y['Close'].iat[-1]; last_x=mdates.date2num(df_1y.index[-1]); xr=ax_d.get_xlim()[1]
    ax_d.add_line(Line2D([last_x,xr],[last_close,last_close],color='yellow',ls='--',lw=1))
    ax_d.text(xr+2,last_close,f"{last_close:.2f}",va='center',ha='left',fontsize=8,color='yellow')

    ax_m.set_title(f"{symbol} 10Y Monthly"); ax_w.set_title(f"{symbol} 10Y Weekly")
    ax_d.set_title(f"{symbol} 1Y Daily with VWAP (YTD: {ytd_str})")

    lines=[Line2D([0],[0],color='lightblue',lw=2),Line2D([0],[0],color='pink',lw=2),Line2D([0],[0],color='yellow',lw=2)]
    labels=[f"High‑anchor VWAP: {df_1y['vwap_maxloc'].iat[-1]:.2f}",f"Low‑anchor VWAP: {df_1y['vwap_minloc'].iat[-1]:.2f}",f"High‑vol VWAP: {df_1y['vwap_highvol'].iat[-1]:.2f}"]
    ax_d.legend(lines,labels,loc='upper left',fontsize=8,framealpha=0.15)

    for a in (ax_m,ax_w,ax_d): _beautify(a)
    fig.tight_layout(); return fig

# ────────────────────────── MAIN ───────────────────────────────────────

def main():
    xls=r"P:\\10_CWP Trade Department\\_Matrix_\\code_outputs\\grip_momo\\grip_allocation.xlsx"
    syms=pd.read_excel(xls,header=2,sheet_name=1)['Ticker'].dropna().tolist()+['QQQ','SPY','IWM','KLAC','TXN','ADBE','MRK','MA','UBER','NOW','GE','LIN']

    out_dir=r"C:\\Users\\rmathews\\Downloads"; os.makedirs(out_dir,exist_ok=True)
    pdf_path,docx_path,zip_path=[os.path.join(out_dir,f) for f in ('vwap_report.pdf','vwap_report.docx','vwap_pngs.zip')]

    doc=Document();sec=doc.sections[0];sec.orientation=WD_ORIENT.LANDSCAPE;sec.page_width,sec.page_height=sec.page_height,sec.page_width
    for side in ('top','bottom','left','right'): setattr(sec,f"{side}_margin",Inches(0.25))
    usable_w=sec.page_width-sec.left_margin-sec.right_margin

    tmp_pngs=[]
    with PdfPages(pdf_path) as pdf:
        for i,s in enumerate(syms,1):
            print(f"{i}/{len(syms)} → {s}"); fig=build_fig(s)
            if fig is None: print('  ↳ skipped'); continue
            pdf.savefig(fig)
            png=os.path.join(out_dir,f"{s}_vwap.png"); fig.savefig(png,dpi=300,bbox_inches='tight',facecolor=fig.get_facecolor()); tmp_pngs.append(png)
            doc.add_picture(png,width=usable_w); doc.add_paragraph(f"{s} — VWAP multi‑time‑frame");
            if i!=len(syms): doc.add_page_break()
            plt.close(fig)

    doc.save(docx_path)

    print('PDF, DOCX saved to:',out_dir)

if __name__=='__main__':
    main()
